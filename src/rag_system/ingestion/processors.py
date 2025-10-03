import PyPDF2
import docx
from pathlib import Path
from typing import Dict, Any
from bs4 import BeautifulSoup
import markdown
import re

from .models import Document, DocumentProcessor


class PDFProcessor(DocumentProcessor):
    """Process PDF documents"""
    
    def can_process(self, file_path: Path) -> bool:
        return file_path.suffix.lower() == '.pdf'
    
    def process(self, file_path: Path) -> Document:
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text_content = []
                
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text.strip():
                            text_content.append(page_text)
                    except Exception as e:
                        print(f"Warning: Could not extract text from page {page_num}: {e}")
                
                content = "\n\n".join(text_content)
                
                # Clean up text
                content = self._clean_text(content)
                
                metadata = {
                    "file_type": "pdf",
                    "file_name": file_path.name,
                    "file_size": file_path.stat().st_size,
                    "num_pages": len(pdf_reader.pages),
                    "page_count": len(text_content)
                }
                
                return Document(
                    id="",  # Will be auto-generated
                    content=content,
                    metadata=metadata,
                    source=str(file_path)
                )
                
        except Exception as e:
            raise Exception(f"Error processing PDF {file_path}: {e}")
    
    def _clean_text(self, text: str) -> str:
        """Clean extracted text"""
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        # Remove page breaks and form feeds
        text = text.replace('\f', '\n').replace('\r', '\n')
        # Normalize line breaks
        text = re.sub(r'\n+', '\n', text)
        return text.strip()


class TextProcessor(DocumentProcessor):
    """Process plain text documents"""
    
    def can_process(self, file_path: Path) -> bool:
        return file_path.suffix.lower() in ['.txt', '.text']
    
    def process(self, file_path: Path) -> Document:
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
            
            metadata = {
                "file_type": "text",
                "file_name": file_path.name,
                "file_size": file_path.stat().st_size,
                "encoding": "utf-8"
            }
            
            return Document(
                id="",  # Will be auto-generated
                content=content,
                metadata=metadata,
                source=str(file_path)
            )
            
        except UnicodeDecodeError:
            # Try with different encoding
            try:
                with open(file_path, 'r', encoding='latin-1') as file:
                    content = file.read()
                
                metadata = {
                    "file_type": "text",
                    "file_name": file_path.name,
                    "file_size": file_path.stat().st_size,
                    "encoding": "latin-1"
                }
                
                return Document(
                    id="",  # Will be auto-generated
                    content=content,
                    metadata=metadata,
                    source=str(file_path)
                )
            except Exception as e:
                raise Exception(f"Error processing text file {file_path}: {e}")


class HTMLProcessor(DocumentProcessor):
    """Process HTML documents"""
    
    def can_process(self, file_path: Path) -> bool:
        return file_path.suffix.lower() in ['.html', '.htm']
    
    def process(self, file_path: Path) -> Document:
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                html_content = file.read()
            
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # Extract title
            title = soup.find('title')
            title_text = title.get_text().strip() if title else ""
            
            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()
            
            # Extract text content
            content = soup.get_text()
            content = self._clean_text(content)
            
            metadata = {
                "file_type": "html",
                "file_name": file_path.name,
                "file_size": file_path.stat().st_size,
                "title": title_text,
                "encoding": "utf-8"
            }
            
            return Document(
                id="",  # Will be auto-generated
                content=content,
                metadata=metadata,
                source=str(file_path)
            )
            
        except Exception as e:
            raise Exception(f"Error processing HTML file {file_path}: {e}")
    
    def _clean_text(self, text: str) -> str:
        """Clean extracted HTML text"""
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        # Normalize line breaks
        text = re.sub(r'\n+', '\n', text)
        return text.strip()


class MarkdownProcessor(DocumentProcessor):
    """Process Markdown documents"""
    
    def can_process(self, file_path: Path) -> bool:
        return file_path.suffix.lower() in ['.md', '.markdown']
    
    def process(self, file_path: Path) -> Document:
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                md_content = file.read()
            
            # Convert markdown to HTML, then extract text
            html = markdown.markdown(md_content)
            soup = BeautifulSoup(html, 'html.parser')
            content = soup.get_text()
            
            # Also keep the original markdown content for metadata
            metadata = {
                "file_type": "markdown",
                "file_name": file_path.name,
                "file_size": file_path.stat().st_size,
                "original_markdown": md_content[:1000] + "..." if len(md_content) > 1000 else md_content,
                "encoding": "utf-8"
            }
            
            return Document(
                id="",  # Will be auto-generated
                content=content,
                metadata=metadata,
                source=str(file_path)
            )
            
        except Exception as e:
            raise Exception(f"Error processing Markdown file {file_path}: {e}")


class DocxProcessor(DocumentProcessor):
    """Process DOCX documents"""
    
    def can_process(self, file_path: Path) -> bool:
        return file_path.suffix.lower() == '.docx'
    
    def process(self, file_path: Path) -> Document:
        try:
            doc = docx.Document(str(file_path))
            
            text_content = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            content = "\n".join(text_content)
            
            # Extract metadata
            core_props = doc.core_properties
            metadata = {
                "file_type": "docx",
                "file_name": file_path.name,
                "file_size": file_path.stat().st_size,
                "title": core_props.title or "",
                "author": core_props.author or "",
                "subject": core_props.subject or "",
                "created": str(core_props.created) if core_props.created else "",
                "modified": str(core_props.modified) if core_props.modified else "",
                "paragraph_count": len(doc.paragraphs)
            }
            
            return Document(
                id="",  # Will be auto-generated
                content=content,
                metadata=metadata,
                source=str(file_path)
            )
            
        except Exception as e:
            raise Exception(f"Error processing DOCX file {file_path}: {e}")


class ProcessorRegistry:
    """Registry for document processors"""
    
    def __init__(self):
        self.processors = [
            PDFProcessor(),
            TextProcessor(),
            HTMLProcessor(),
            MarkdownProcessor(),
            DocxProcessor()
        ]
    
    def get_processor(self, file_path: Path) -> DocumentProcessor:
        """Get appropriate processor for file"""
        for processor in self.processors:
            if processor.can_process(file_path):
                return processor
        
        raise ValueError(f"No processor found for file type: {file_path.suffix}")
    
    def add_processor(self, processor: DocumentProcessor) -> None:
        """Add a custom processor"""
        self.processors.append(processor)